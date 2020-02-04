# -*- coding=utf-8 -*-

import jieba as jb
# import docx
from docx import Document
from deepdiff import DeepDiff
import difflib
import webbrowser


# text = "我开始分析文本"
# seg_list1 = jb.cut(text, cut_all=True)
# print("Full Mode: " + "/".join(seg_list1))
#
# seg_list2 = jb.cut(text, cut_all=False)
# print("Default Mode: " + "/".join(seg_list2))
#
# seg_list3 = jb.cut(text)
# print("Default Mode: " + "/".join(seg_list2))

doc_path1 = "E:\\行业研究\\NLP\\集合资产管理计划合同指引.docx"
doc_path2 = "E:\\行业研究\\NLP\\华菁资管迅投FOF1号集合资产管理计划资产管理合同(1).docx"

doc_model = Document(doc_path1)
doc_target = Document(doc_path2)

diff = difflib.HtmlDiff()
diff2 = difflib.Differ()

plist = []
qlist = []
for p, q in zip(doc_model.paragraphs, doc_target.paragraphs):
    plist.append(p.text)
    qlist.append(q.text)

result = diff.make_file(plist, qlist)
file = open("E:\\行业研究\\NLP\\CompareResult.html", 'w', encoding="utf-8")
file.write(result)
file.close()
webbrowser.open("E:\\行业研究\\NLP\\CompareResult.html")

result2 = diff2.compare(plist, qlist)
print("\n".join(list(result2)))

